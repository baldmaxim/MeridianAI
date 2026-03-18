"""Document loader for meeting context (PDF, MD, TXT, DOCX, XLSX)."""

import io
import logging
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import List, Optional, Tuple, Union

try:
    from PyPDF2 import PdfReader
    PYPDF2_AVAILABLE = True
except ImportError:
    PYPDF2_AVAILABLE = False

try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    from openpyxl import load_workbook
    OPENPYXL_AVAILABLE = True
except ImportError:
    OPENPYXL_AVAILABLE = False

from .document_chunker import DocumentChunker
from .document_retriever import DocumentRetriever

logger = logging.getLogger("ai_helper.document_loader")

SUPPORTED_EXTENSIONS = {".pdf", ".md", ".txt", ".docx", ".xlsx"}

DOC_TYPE_LABELS = {
    "contract": "Договор",
    "bor": "ВОР",
    "estimate": "Смета",
    "specification": "Спецификация",
    "other": "Другое",
}


@dataclass
class MeetingDocument:
    """Loaded meeting document."""
    filename: str
    content: str
    doc_type: str
    loaded_at: datetime
    page_count: int


# --------------- extraction per format ---------------

def _extract_pdf(source: Union[Path, bytes]) -> Tuple[str, int]:
    if not PYPDF2_AVAILABLE:
        raise RuntimeError("PyPDF2 not installed")
    if isinstance(source, bytes):
        reader = PdfReader(io.BytesIO(source))
    else:
        reader = PdfReader(str(source))
    pages_text = []
    for page in reader.pages:
        text = page.extract_text()
        if text:
            pages_text.append(text.strip())
    return "\n\n".join(pages_text), len(reader.pages)


def _extract_md(source: Union[Path, bytes]) -> Tuple[str, int]:
    if isinstance(source, bytes):
        return source.decode("utf-8", errors="replace"), 1
    return Path(source).read_text(encoding="utf-8", errors="replace"), 1


def _extract_txt(source: Union[Path, bytes]) -> Tuple[str, int]:
    if isinstance(source, bytes):
        return source.decode("utf-8", errors="replace"), 1
    return Path(source).read_text(encoding="utf-8", errors="replace"), 1


def _extract_docx(source: Union[Path, bytes]) -> Tuple[str, int]:
    if not DOCX_AVAILABLE:
        raise RuntimeError("python-docx not installed")
    if isinstance(source, bytes):
        doc = DocxDocument(io.BytesIO(source))
    else:
        doc = DocxDocument(str(source))
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    return "\n\n".join(paragraphs), 1


def _extract_xlsx(source: Union[Path, bytes]) -> Tuple[str, int]:
    if not OPENPYXL_AVAILABLE:
        raise RuntimeError("openpyxl not installed")
    if isinstance(source, bytes):
        wb = load_workbook(io.BytesIO(source), read_only=True, data_only=True)
    else:
        wb = load_workbook(str(source), read_only=True, data_only=True)
    parts = []
    for sheet in wb.sheetnames:
        ws = wb[sheet]
        rows = []
        for row in ws.iter_rows(values_only=True):
            cells = [str(c) if c is not None else "" for c in row]
            line = "\t".join(cells).strip()
            if line:
                rows.append(line)
        if rows:
            parts.append(f"[Лист: {sheet}]\n" + "\n".join(rows))
    wb.close()
    return "\n\n".join(parts), len(wb.sheetnames)


def _extract_content(filename: str, source: Union[Path, bytes]) -> Tuple[str, int]:
    ext = Path(filename).suffix.lower()
    if ext == ".pdf":
        return _extract_pdf(source)
    elif ext == ".md":
        return _extract_md(source)
    elif ext == ".txt":
        return _extract_txt(source)
    elif ext == ".docx":
        return _extract_docx(source)
    elif ext == ".xlsx":
        return _extract_xlsx(source)
    else:
        raise ValueError(
            f"Формат {ext} не поддерживается. "
            f"Поддерживаемые: {', '.join(sorted(SUPPORTED_EXTENSIONS))}"
        )


# --------------- main class ---------------

class DocumentLoader:
    """Loads and manages documents for meeting context."""

    def __init__(self):
        self.documents: List[MeetingDocument] = []
        self.meeting_topic: str = ""
        self.meeting_notes: str = ""
        self._chunker = DocumentChunker()
        self._retriever = DocumentRetriever()

    def load_file(self, path, doc_type: str = "other") -> Optional[MeetingDocument]:
        """Load a file and extract text."""
        path = Path(path)
        if not path.exists():
            raise FileNotFoundError(f"File not found: {path}")

        ext = path.suffix.lower()
        if ext not in SUPPORTED_EXTENSIONS:
            raise ValueError(
                f"Формат {ext} не поддерживается. "
                f"Поддерживаемые: {', '.join(sorted(SUPPORTED_EXTENSIONS))}"
            )

        content, page_count = _extract_content(path.name, path)
        if not content.strip():
            raise ValueError(
                f"Не удалось извлечь текст из {path.name}. "
                "Файл может быть пустым или содержать только изображения."
            )

        doc = MeetingDocument(
            filename=path.name,
            content=content,
            doc_type=doc_type,
            loaded_at=datetime.now(),
            page_count=page_count,
        )
        self.documents.append(doc)
        self._index_doc(doc)
        return doc

    def load_bytes(self, filename: str, content_bytes: bytes,
                   doc_type: str = "other") -> MeetingDocument:
        """Load a document from bytes (for web upload)."""
        ext = Path(filename).suffix.lower()
        if ext not in SUPPORTED_EXTENSIONS:
            raise ValueError(
                f"Формат {ext} не поддерживается. "
                f"Поддерживаемые: {', '.join(sorted(SUPPORTED_EXTENSIONS))}"
            )

        content, page_count = _extract_content(filename, content_bytes)
        if not content.strip():
            raise ValueError(
                f"Не удалось извлечь текст из {filename}. "
                "Файл может быть пустым или содержать только изображения."
            )

        doc = MeetingDocument(
            filename=filename,
            content=content,
            doc_type=doc_type,
            loaded_at=datetime.now(),
            page_count=page_count,
        )
        self.documents.append(doc)
        self._index_doc(doc)
        return doc

    def _index_doc(self, doc: MeetingDocument) -> None:
        """Chunk and index a document for retrieval."""
        chunks = self._chunker.chunk(doc)
        self._retriever.index(chunks)

    def remove_document(self, filename: str):
        self.documents = [d for d in self.documents if d.filename != filename]
        self._retriever.remove(filename)

    def clear(self):
        self.documents.clear()
        self.meeting_topic = ""
        self.meeting_notes = ""
        self._retriever.clear()

    def retrieve_relevant(self, query: str, max_chunks: int = 5,
                          max_chars: int = 2000) -> str:
        """Retrieve document chunks relevant to query. Returns formatted text or ''."""
        return self._retriever.retrieve(query, max_chunks=max_chunks, max_chars=max_chars)

    def get_context_for_prompt(self, max_chars: int = 3000) -> str:
        """Get combined document context for LLM prompt, with smart truncation."""
        parts = []
        if self.meeting_topic:
            parts.append(f"ТЕМА ВСТРЕЧИ: {self.meeting_topic}")
        if self.meeting_notes:
            parts.append(f"ДОПОЛНИТЕЛЬНЫЙ КОНТЕКСТ:\n{self.meeting_notes}")
        if self.documents:
            chars_per_doc = max_chars // len(self.documents)
            doc_parts = []
            for doc in self.documents:
                label = DOC_TYPE_LABELS.get(doc.doc_type, doc.doc_type)
                ext = Path(doc.filename).suffix.lower()
                count_label = f"{doc.page_count} стр." if ext == ".pdf" else ""
                meta = f" ({count_label})" if count_label else ""
                header = f"--- {label}: {doc.filename}{meta} ---"
                if len(doc.content) <= chars_per_doc:
                    doc_parts.append(f"{header}\n{doc.content}")
                else:
                    half = chars_per_doc // 2
                    truncated = doc.content[:half] + "\n[...]\n" + doc.content[-half:]
                    doc_parts.append(f"{header}\n{truncated}")
            parts.append("ДОКУМЕНТЫ ВСТРЕЧИ:\n" + "\n\n".join(doc_parts))
        return "\n\n".join(parts)

    def get_document_context(self, max_chars: int = 3000) -> str:
        """Get only document content for prompt (without topic/notes)."""
        if not self.documents:
            return ""
        chars_per_doc = max_chars // len(self.documents)
        doc_parts = []
        for doc in self.documents:
            label = DOC_TYPE_LABELS.get(doc.doc_type, doc.doc_type)
            ext = Path(doc.filename).suffix.lower()
            if ext == ".xlsx":
                count_label = f"{doc.page_count} лист."
            elif ext in (".txt", ".md"):
                count_label = ""
            else:
                count_label = f"{doc.page_count} стр."
            meta = f" ({count_label})" if count_label else ""
            header = f"--- {label}: {doc.filename}{meta} ---"
            if len(doc.content) <= chars_per_doc:
                doc_parts.append(f"{header}\n{doc.content}")
            else:
                half = chars_per_doc // 2
                truncated = doc.content[:half] + "\n[...]\n" + doc.content[-half:]
                doc_parts.append(f"{header}\n{truncated}")
        return "ДОКУМЕНТЫ ВСТРЕЧИ:\n" + "\n\n".join(doc_parts)

    def has_context(self) -> bool:
        return bool(self.documents or self.meeting_topic or self.meeting_notes)
